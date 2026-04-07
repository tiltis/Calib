"""piza2 (v9) vs piza3 (v10) 비교 문서 생성"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = r"C:\Users\tilti\OneDrive\samjung\output\v9_vs_v10_비교.docx"

doc = Document()

# ── 스타일 설정 ──
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.25

# ── 제목 ──
title = doc.add_heading("파도 영상 왜곡 보정 알고리즘 비교", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("piza2.py (v9)  vs  piza3.py (v10)")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()  # 여백

# ── 1. 개요 ──
doc.add_heading("1. 개요", level=1)
doc.add_paragraph(
    "본 문서는 해상 파도 영상의 등간격(m/pixel) 변환을 수행하는 "
    "piza2.py(v9)와 piza3.py(v10)의 핵심 차이를 정리한다."
)

# ── 2. 버전 요약 표 ──
doc.add_heading("2. 버전 요약", level=1)

table = doc.add_table(rows=7, cols=3, style="Light Grid Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["항목", "v9 (piza2.py)", "v10 (piza3.py)"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].bold = True

rows_data = [
    ["보정 방향", "세로(Y)만", "세로(Y) + 가로(X) 동시"],
    ["build_remap_tables\nmap_x 계산",
     "map_x = col\n(가로 좌표 그대로 유지)",
     "map_x = cx + (col - cx) / scale_x\n(행별 수평 스케일 적용)"],
    ["방사형 왜곡 처리", "미처리\n→ 파도가 방사형으로 잔존",
     "제거\n→ 먼 거리 파도 마루를 물리적 너비로 복원"],
    ["weather_cache.json 탐색", "현재 폴더만 탐색",
     "현재 폴더 + 상위 폴더 탐색\n(하위 폴더 실행 대응)"],
    ["파일 위치", "parang/piza2.py\n(수정 금지 — 안정 기준선)",
     "parang/2026-04-06/piza3.py\n(활성 개발본)"],
    ["기타 로직", "동일 (CrestTracker, velocity 추정, R² 판단 등)", "동일"],
]

for r, row_data in enumerate(rows_data):
    for c, val in enumerate(row_data):
        table.rows[r + 1].cells[c].text = val

# ── 3. 핵심 변경: XY 동시 보정 ──
doc.add_heading("3. 핵심 변경: XY 동시 보정", level=1)

doc.add_heading("3.1 문제 (v9)", level=2)
doc.add_paragraph(
    "v9의 build_remap_tables()는 세로 방향(Y)만 비선형 보정을 수행하고, "
    "가로 방향(X)은 map_x = col 그대로 유지했다. "
    "이로 인해 카메라 원근에 의한 수평 방향 축소가 보정되지 않아, "
    "결과 영상에서 파도가 여전히 방사형(부채꼴)으로 보이는 문제가 있었다."
)

doc.add_heading("3.2 해결 (v10)", level=2)
doc.add_paragraph(
    "v10에서는 각 출력 행 r에 대해 해당 소스 행의 m/pixel 값으로부터 "
    "수평 스케일 팩터 scale_x를 계산한다:"
)

# 수식 블록
eq = doc.add_paragraph()
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq.paragraph_format.space_before = Pt(8)
eq.paragraph_format.space_after = Pt(8)
run = eq.add_run("scale_x(r) = mpp(src_y[r]) / base_mpp")
run.font.name = "Consolas"
run.font.size = Pt(11)
run.bold = True

eq2 = doc.add_paragraph()
eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq2.paragraph_format.space_after = Pt(12)
run2 = eq2.add_run("map_x[r, col] = cx + (col − cx) / scale_x(r)")
run2.font.name = "Consolas"
run2.font.size = Pt(11)
run2.bold = True

doc.add_paragraph(
    "상단(먼 거리)의 행은 scale_x > 1이므로, 소스의 중앙 부근을 넓게 펼쳐서 "
    "물리적 파도 마루 폭을 복원한다. "
    "하단(기준 행)은 scale_x ≈ 1로 변화 없다."
)

# ── 3.3 동작 비교 표 ──
doc.add_heading("3.3 행별 동작 비교", level=2)
t2 = doc.add_table(rows=4, cols=4, style="Light Grid Accent 1")
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["영역", "위치", "v9 map_x", "v10 map_x"]):
    t2.rows[0].cells[i].text = h
    for p in t2.rows[0].cells[i].paragraphs:
        p.runs[0].bold = True

t2_data = [
    ["상단 (먼 거리)", "m/pixel 큼", "col (그대로)", "중앙으로부터 펼침 (확대)"],
    ["중간", "m/pixel 중간", "col (그대로)", "약간 펼침"],
    ["하단 (가까운 거리)", "base_mpp", "col (그대로)", "col (변화 없음)"],
]
for r, row_data in enumerate(t2_data):
    for c, val in enumerate(row_data):
        t2.rows[r + 1].cells[c].text = val

# ── 4. 부가 변경 ──
doc.add_heading("4. 부가 변경", level=1)

doc.add_heading("4.1 weather_cache.json 탐색 경로", level=2)
doc.add_paragraph(
    "v10은 piza3.py가 하위 폴더(2026-04-06/)에 위치하므로, "
    "weather_cache.json을 현재 폴더뿐 아니라 상위 폴더에서도 탐색하도록 변경했다."
)

# v9 코드
doc.add_paragraph("v9:", style="List Bullet")
p_v9 = doc.add_paragraph()
run_v9 = p_v9.add_run(
    'cache_path = os.path.join(os.path.dirname(__file__), "weather_cache.json")'
)
run_v9.font.name = "Consolas"
run_v9.font.size = Pt(9)

# v10 코드
doc.add_paragraph("v10:", style="List Bullet")
p_v10 = doc.add_paragraph()
run_v10 = p_v10.add_run(
    "for _d in [_script_dir, os.path.dirname(_script_dir)]:\n"
    '    _c = os.path.join(_d, "weather_cache.json")\n'
    "    if os.path.exists(_c): cache_path = _c; break"
)
run_v10.font.name = "Consolas"
run_v10.font.size = Pt(9)

# ── 5. 동일한 부분 ──
doc.add_heading("5. 변경 없는 부분", level=1)
unchanged = [
    "CrestTracker: Sobel Y 기반 파봉 추적",
    "get_velocity_vs_y(): 슬라이딩 윈도우 속도 추정",
    "build_velocity_scale_map(): 선형회귀 m/pixel(y) → R² 판단",
    "encounter speed 선택: swell dominance 기반 가중 enc speed",
    "output 저장: calib.json + ortho.mp4",
    "ROI 선택 UI, 캘리브레이션 Phase 1/2 흐름",
]
for item in unchanged:
    doc.add_paragraph(item, style="List Bullet")

# ── 6. 요약 ──
doc.add_heading("6. 요약", level=1)
doc.add_paragraph(
    "v10의 유일한 핵심 변경은 build_remap_tables()에서 "
    "가로(X) 방향 원근 보정을 추가한 것이다. "
    "이를 통해 v9에서 남아있던 방사형 파도 형태가 제거되고, "
    "상단(먼 거리)의 파도 마루가 물리적 실제 폭으로 펼쳐져 "
    "등방성(isotropic) 등간격 투영이 달성된다."
)

# ── 저장 ──
doc.save(OUT)
print(f"저장 완료: {OUT}")
